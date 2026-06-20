import io
from enum import Enum

import fitz  # PyMuPDF
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document

from log_config.logger_config import logger


# =========================
# LOADERS
# =========================

class PDFLoader:
    def load(self, data: bytes) -> str:
        try:
            logger.debug("Loading PDF document...")
            doc = fitz.open(stream=data, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)

            logger.debug(
                f"PDF loaded successfully. Pages: {len(doc)}, Characters: {len(text)}"
            )
            return text

        except Exception as e:
            logger.exception(f"Error while loading PDF: {e}")
            raise


class DOCXLoader:
    def load(self, data: bytes) -> str:
        try:
            logger.debug("Loading DOCX document...")

            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)

            logger.debug(
                f"DOCX loaded successfully. Paragraphs: {len(doc.paragraphs)}, Characters: {len(text)}"
            )
            return text

        except Exception as e:
            logger.exception(f"Error while loading DOCX: {e}")
            raise


class CSVLoader:
    def load(self, data: bytes) -> str:
        try:
            logger.debug("Loading CSV file...")

            df = pd.read_csv(io.BytesIO(data))
            text = df.to_string(index=False)

            logger.debug(
                f"CSV loaded successfully. Rows: {len(df)}, Columns: {len(df.columns)}"
            )
            return text

        except Exception as e:
            logger.exception(f"Error while loading CSV: {e}")
            raise


class HTMLLoader:
    def load(self, data: bytes) -> str:
        try:
            logger.debug("Loading HTML document...")

            soup = BeautifulSoup(data, "lxml")
            text = soup.get_text(separator="\n", strip=True)

            logger.debug(
                f"HTML loaded successfully. Characters extracted: {len(text)}"
            )
            return text

        except Exception as e:
            logger.exception(f"Error while loading HTML: {e}")
            raise


class TextLoader:
    def load(self, data: bytes) -> str:
        try:
            logger.debug("Loading text file...")

            text = data.decode("utf-8", errors="ignore")

            logger.debug(
                f"Text file loaded successfully. Characters: {len(text)}"
            )
            return text

        except Exception as e:
            logger.exception(f"Error while loading text file: {e}")
            raise