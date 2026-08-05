import io

import fitz  # PyMuPDF
import pandas as pd
from bs4 import BeautifulSoup
from core.logger_config import logger
from docx import Document

# =========================
# LOADERS
# =========================


class PDFLoader:
    """
    Loader for PDF documents.

    Uses PyMuPDF to extract text content from PDF files.
    """

    def load(self, data: bytes) -> str:
        """
        Extract text from PDF binary data.

        Args:
            data (bytes):
                Raw PDF file content.

        Returns:
            str:
                Extracted text from all PDF pages.

        Raises:
            Exception:
                If PDF parsing fails.

        Example:
            >>> loader = PDFLoader()
            >>> text = loader.load(pdf_bytes)
        """

        try:
            logger.debug("Loading PDF document...")

            doc = fitz.open(stream=data, filetype="pdf")

            text = "\n".join(page.get_text() for page in doc)

            logger.debug(f"PDF loaded successfully. Pages: {len(doc)}, Characters: {len(text)}")

            return text

        except Exception as e:
            logger.exception(f"Error while loading PDF: {e}")
            raise


class DOCXLoader:
    """
    Loader for Microsoft Word DOCX documents.

    Extracts text from paragraphs using python-docx.
    """

    def load(self, data: bytes) -> str:
        """
        Extract text from DOCX binary content.

        Args:
            data (bytes):
                Raw DOCX file bytes.

        Returns:
            str:
                Text extracted from document paragraphs.

        Raises:
            Exception:
                If DOCX parsing fails.
        """
        try:
            logger.debug("Loading DOCX document...")

            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)

            logger.debug(
                f"DOCX loaded successfully. Paragraphs: {len(doc.paragraphs)}, Characters: {len(text)}"
            )
            return text

        except Exception as e:
            logger.exception(f"Error while loading DOCX: {e}")
            raise


class CSVLoader:
    """
    Loader for CSV files.

    Converts tabular CSV data into textual representation
    suitable for downstream processing and embeddings.
    """

    def load(self, data: bytes) -> str:
        """
        Convert CSV content into text.

        Args:
            data (bytes):
                Raw CSV file bytes.

        Returns:
            str:
                String representation of CSV table.

        Raises:
            Exception:
                If CSV parsing fails.
        """
        try:
            logger.debug("Loading CSV file...")

            df = pd.read_csv(io.BytesIO(data))
            text = df.to_string(index=False)

            logger.debug(f"CSV loaded successfully. Rows: {len(df)}, Columns: {len(df.columns)}")
            return text

        except Exception as e:
            logger.exception(f"Error while loading CSV: {e}")
            raise


class HTMLLoader:
    """
    Loader for HTML documents.

    Removes HTML tags and extracts readable text content.
    """

    def load(self, data: bytes) -> str:
        """
        Extract visible text from HTML.

        Args:
            data (bytes):
                Raw HTML content.

        Returns:
            str:
                Clean text without HTML tags.

        Raises:
            Exception:
                If HTML parsing fails.
        """

    def load(self, data: bytes) -> str:
        try:
            logger.debug("Loading HTML document...")

            soup = BeautifulSoup(data, "lxml")
            text = soup.get_text(separator="\n", strip=True)

            logger.debug(f"HTML loaded successfully. Characters extracted: {len(text)}")
            return text

        except Exception as e:
            logger.exception(f"Error while loading HTML: {e}")
            raise


class TextLoader:
    """
    Loader for plain text files.

    Decodes UTF-8 encoded files into string format.
    """

    def load(self, data: bytes) -> str:
        """
        Decode text file bytes.

        Args:
            data (bytes):
                Raw text file content.

        Returns:
            str:
                Decoded UTF-8 text.

        Raises:
            Exception:
                If decoding fails.
        """
        try:
            logger.debug("Loading text file...")

            text = data.decode("utf-8", errors="ignore")

            logger.debug(f"Text file loaded successfully. Characters: {len(text)}")
            return text

        except Exception as e:
            logger.exception(f"Error while loading text file: {e}")
            raise
